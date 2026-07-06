from __future__ import annotations

from pathlib import Path
import argparse

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


FEATURES = [
    ("FR-001", "任务理解", "多轮需求澄清与合同范围锁定", "P0", "用户提交目标或项目文档后，智能体能识别目标、约束、交付物、验收口径，并生成待确认问题清单。"),
    ("FR-002", "计划编排", "执行计划生成与阶段门禁", "P0", "系统能把需求拆成探索、设计、实现、验证、交付阶段，并标记每阶段输入、输出、负责人和退出条件。"),
    ("FR-003", "代码理解", "仓库扫描与上下文索引", "P0", "智能体能读取项目结构、关键文件、依赖和测试入口，形成可追溯上下文包。"),
    ("FR-004", "代码修改", "受控文件编辑与补丁生成", "P0", "智能体只能在授权目录内修改文件，修改前后能说明影响范围并生成可审计补丁。"),
    ("FR-005", "测试验证", "自动运行测试与错误归因", "P0", "智能体能选择相关测试命令，汇总失败日志，定位到文件、函数或配置层面的原因。"),
    ("FR-006", "文档生成", "需求、设计、变更和验收文档生成", "P1", "系统能输出 Word、PDF、Markdown、Excel 等交付文档，并记录版本、来源和生成时间。"),
    ("FR-007", "版本控制", "Git 分支、提交和变更摘要", "P1", "智能体能在用户授权后创建分支、提交变更、生成 PR 摘要，并避免覆盖用户未授权改动。"),
    ("FR-008", "安全权限", "高风险动作拦截与确认", "P0", "删除、移动、外部网络访问、执行脚本等动作需要风险等级和确认策略。"),
    ("FR-009", "工具调用", "Shell、浏览器、文件和模型工具编排", "P1", "系统能按任务选择工具，保存工具输入输出摘要，失败时给出重试策略。"),
    ("FR-010", "模型路由", "多模型配置与降级", "P1", "支持按任务选择不同模型，模型不可用时降级到备用配置或本地 Stub。"),
    ("FR-011", "知识沉淀", "项目记忆和经验复用", "P2", "关键决策、需求、接口约束和验收记录可进入项目记忆供后续任务检索。"),
    ("FR-012", "交互体验", "前端实时状态、预览和产物入口", "P1", "用户能看到来源、解析预览、执行步骤、输出文件和下一步动作。"),
    ("FR-013", "异常恢复", "中断续跑与失败回滚建议", "P1", "任务中断后可恢复上下文，失败时展示失败阶段、原因和可执行修复建议。"),
    ("FR-014", "审计日志", "全链路追踪和 TraceID", "P0", "每次执行生成 TraceID，记录输入来源、模型、工具、文件产物和运行结果。"),
    ("FR-015", "集成接口", "开放 API 与桌面桥接", "P1", "提供稳定的本地 API 和 Electron IPC 契约，支持前端、自动化和测试调用。"),
    ("FR-016", "验收助手", "合同条款到验收用例映射", "P1", "系统能把功能项映射为验收标准、测试建议和交付检查清单。"),
]

NFRS = [
    ("NFR-001", "可靠性", "核心链路失败后应返回可理解错误、TraceID 和用户可执行恢复动作。"),
    ("NFR-002", "安全性", "默认限制文件访问在项目目录内，不保存 API Key 明文到产物。"),
    ("NFR-003", "可维护性", "后端模块边界清晰，外部 SDK 封装在 adapter 层，应用层依赖 ports。"),
    ("NFR-004", "性能", "普通项目的需求解析与功能清单生成应在 3 分钟内完成，长文档可进入异步队列。"),
    ("NFR-005", "兼容性", "生成的 DOCX、XLSX、PDF 应能在 Microsoft Office、WPS 和常见 PDF 阅读器打开。"),
]

MILESTONES = [
    ("M1", "合同需求确认", "完成需求文档解析、功能清单、需求规格说明初版。"),
    ("M2", "编码智能体 MVP", "完成仓库理解、计划生成、补丁编辑、测试验证闭环。"),
    ("M3", "安全与审计", "完成权限门禁、TraceID、变更审计和用户确认策略。"),
    ("M4", "交付验收", "输出验收报告、部署说明、用户操作手册和回归测试记录。"),
]

CLAUSES = [
    ("1. 项目背景与合同目标", [
        "甲方拟建设一套面向软件工程项目的编码智能体，用于把自然语言需求、项目文档和代码仓库上下文转化为可执行工程变更。",
        "乙方交付的系统应覆盖任务理解、计划编排、代码修改、测试验证、文档产物和审计追踪等能力。",
        "本文件作为合同式需求基线，用于后续功能拆分、报价、里程碑验收和变更控制。",
    ]),
    ("2. 合同范围", [
        "范围内：桌面工作台、Main Runtime、本地项目空间、模型路由、工具调用、文档生成、Trace 审计。",
        "范围外：替代企业代码评审制度、绕过权限审批的自动上线、未授权读取项目目录外文件。",
        "任何新增范围应通过变更单记录目标、成本、工期、风险和验收标准。",
    ]),
    ("3. 角色与责任", [
        "甲方产品负责人负责确认业务目标、验收标准、数据边界和上线策略。",
        "甲方技术负责人负责提供仓库、测试命令、部署约束和代码规范。",
        "乙方负责需求分析、系统设计、实现、测试、文档和交付培训。",
    ]),
    ("4. 验收原则", [
        "每个功能项必须具备功能编号、用户角色、触发场景、输入、输出、优先级和验收标准。",
        "验收以可运行演示、生成文档、测试结果和审计记录为准。",
        "若第三方模型或外部服务不可用，应验证降级策略和错误提示。",
    ]),
]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project-root", default=r"C:\project\DreamWorkerTest")
    args = parser.parse_args()
    out_dir = Path(args.project_root) / "workspace" / "imports" / "requirements" / "test-fixtures"
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "coding-agent-contract-requirements.docx"
    pdf_path = out_dir / "coding-agent-contract-requirements.pdf"
    build_docx(docx_path)
    build_pdf(pdf_path)
    print(docx_path)
    print(pdf_path)


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 16, 8),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 12, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 8, 4),
    ]:
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = color
        styles[style_name].paragraph_format.space_before = Pt(before)
        styles[style_name].paragraph_format.space_after = Pt(after)
        styles[style_name].paragraph_format.line_spacing = 1.10
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "DreamWorker 编码智能体需求基线"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    footer = section.footer.paragraphs[0]
    footer.text = "Confidential - Contract Requirements Baseline"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("编码智能体完整合同式需求文档")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle = doc.add_paragraph("测试用例版本：Word/PDF 双版 | 适用模块：DreamWorker 需求分析 | 生成日期：2026-07-05")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_docx_table(doc, ["字段", "内容"], [
        ("文档编号", "DW-REQ-CA-20260705"),
        ("合同甲方", "DreamWorker 项目方"),
        ("合同乙方", "编码智能体交付方"),
        ("项目名称", "编码智能体 Coding Agent"),
        ("需求基线", "v1.0 - 用于上传解析、功能清单生成和需求规格说明生成的验收测试"),
    ])

    for title_text, bullets in CLAUSES:
        doc.add_heading(title_text, level=1)
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("5. 功能需求清单", level=1)
    add_docx_table(doc, ["ID", "模块", "功能", "优先级", "验收摘要"], FEATURES)
    doc.add_heading("6. 非功能需求", level=1)
    add_docx_table(doc, ["ID", "类别", "要求"], NFRS)
    doc.add_heading("7. 里程碑与交付物", level=1)
    add_docx_table(doc, ["里程碑", "名称", "交付/验收说明"], MILESTONES)
    doc.add_heading("8. 交付产物清单", level=1)
    for item in [
        "功能清单 Excel：包含功能 ID、模块、名称、角色、场景、优先级、验收标准、依赖和来源。",
        "需求规格说明 Word：包含项目背景、需求来源、角色、功能需求、非功能需求、风险和待确认问题。",
        "结构化 JSON：供后续 PRD、原型、蓝图和开发计划模块读取。",
        "验收记录：包含上传文件、MinerU 解析预览、运行 TraceID、输出文件路径和渲染检查结果。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("9. 风险、假设与变更控制", level=1)
    for item in [
        "模型输出可能存在不完整或幻觉，需要通过来源引用和验收清单约束。",
        "第三方 MinerU Open API 或本地 CLI 不可用时，系统必须返回可恢复错误。",
        "所有范围变更需由甲方产品负责人确认，并形成变更记录。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("10. 签署与生效", level=1)
    add_docx_table(doc, ["签署方", "代表", "日期", "备注"], [
        ("甲方", "项目负责人", "2026-07-05", "确认需求基线和验收口径"),
        ("乙方", "交付负责人", "2026-07-05", "承诺按本文档交付并接受验收"),
    ])
    doc.save(path)


def add_docx_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        set_docx_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_docx_cell_text(cell, value)
    doc.add_paragraph("")


def set_docx_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_pdf(path: Path) -> None:
    pdfmetrics.registerFont(TTFont("SimHei", r"C:\Windows\Fonts\simhei.ttf"))
    base = getSampleStyleSheet()
    body = ParagraphStyle("CNBody", parent=base["Normal"], fontName="SimHei", fontSize=9.2, leading=13.2, spaceAfter=6)
    small = ParagraphStyle("CNSmall", parent=body, fontSize=8, leading=11, textColor=colors.HexColor("#555555"))
    h1 = ParagraphStyle("CNH1", parent=body, fontSize=14, leading=18, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=6)
    title = ParagraphStyle("CNTitle", parent=body, fontSize=19, leading=24, alignment=TA_CENTER, textColor=colors.HexColor("#0B2545"), spaceAfter=8)

    story = [
        Paragraph(escape("编码智能体完整合同式需求文档"), title),
        Paragraph(escape("测试用例版本：Word/PDF 双版 | 适用模块：DreamWorker 需求分析 | 生成日期：2026-07-05"), small),
        Spacer(1, 8),
    ]
    story.append(pdf_table(["字段", "内容"], [
        ("文档编号", "DW-REQ-CA-20260705"),
        ("合同甲方", "DreamWorker 项目方"),
        ("合同乙方", "编码智能体交付方"),
        ("项目名称", "编码智能体 Coding Agent"),
        ("需求基线", "v1.0 - 用于上传解析、功能清单生成和需求规格说明生成的验收测试"),
    ], [1.4, 5.0], small))
    story.append(Spacer(1, 8))
    for title_text, bullets in CLAUSES:
        story.append(Paragraph(escape(title_text), h1))
        for bullet in bullets:
            story.append(Paragraph(escape("• " + bullet), body))
    story.append(Paragraph(escape("5. 功能需求清单"), h1))
    story.append(pdf_table(["ID", "模块", "功能", "优先级", "验收摘要"], FEATURES, [0.55, 0.75, 1.15, 0.45, 3.55], small))
    story.append(Spacer(1, 8))
    story.append(Paragraph(escape("6. 非功能需求"), h1))
    story.append(pdf_table(["ID", "类别", "要求"], NFRS, [0.75, 0.9, 4.75], small))
    story.append(Spacer(1, 8))
    story.append(Paragraph(escape("7. 里程碑与交付物"), h1))
    story.append(pdf_table(["里程碑", "名称", "交付/验收说明"], MILESTONES, [0.7, 1.25, 4.45], small))
    story.append(Paragraph(escape("8. 交付产物清单"), h1))
    for item in [
        "功能清单 Excel：包含功能 ID、模块、名称、角色、场景、优先级、验收标准、依赖和来源。",
        "需求规格说明 Word：包含项目背景、需求来源、角色、功能需求、非功能需求、风险和待确认问题。",
        "结构化 JSON：供后续 PRD、原型、蓝图和开发计划模块读取。",
        "验收记录：包含上传文件、MinerU 解析预览、运行 TraceID、输出文件路径和渲染检查结果。",
    ]:
        story.append(Paragraph(escape("• " + item), body))
    story.append(Paragraph(escape("9. 风险、假设与变更控制"), h1))
    for item in [
        "模型输出可能存在不完整或幻觉，需要通过来源引用和验收清单约束。",
        "第三方 MinerU Open API 或本地 CLI 不可用时，系统必须返回可恢复错误。",
        "所有范围变更需由甲方产品负责人确认，并形成变更记录。",
    ]:
        story.append(Paragraph(escape("• " + item), body))
    story.append(Paragraph(escape("10. 签署与生效"), h1))
    story.append(pdf_table(["签署方", "代表", "日期", "备注"], [
        ("甲方", "项目负责人", "2026-07-05", "确认需求基线和验收口径"),
        ("乙方", "交付负责人", "2026-07-05", "承诺按本文档交付并接受验收"),
    ], [1.0, 1.3, 1.0, 3.1], small))

    def footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont("SimHei", 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(inch, 0.55 * inch, "DreamWorker 编码智能体需求基线")
        canvas.drawRightString(letter[0] - inch, 0.55 * inch, f"Page {doc_obj.page}")
        canvas.restoreState()

    pdf = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    pdf.build(story, onFirstPage=footer, onLaterPages=footer)


def pdf_table(headers: list[str], rows: list[tuple[str, ...]], widths: list[float], style: ParagraphStyle) -> Table:
    data = [[Paragraph(escape(header), style) for header in headers]]
    data.extend([[Paragraph(escape(value), style) for value in row] for row in rows])
    table = Table(data, colWidths=[width * inch for width in widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "SimHei"),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def escape(value: str) -> str:
    return str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


if __name__ == "__main__":
    main()
